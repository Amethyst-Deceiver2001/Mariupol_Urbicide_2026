#!/usr/bin/env python3
"""Content-level triage of every document attachment pulled by scripts/233
(telegram_document_media, 509 files across 57 channels as of 2026-07-06) —
decrees, ownerless-property lists, government press releases, resident
complaint letters, etc. Prior surveys (scripts/247-249/258/263) only ever
keyword-matched Telegram MESSAGE TEXT; this is the first pass that opens the
actual PDF/DOCX/XLSX attachments and extracts their content, so a decree or
list buried inside a document with an innocuous caption doesn't get missed.

Local, offline, no network — reads each file straight from data/raw/ via its
sha256, matches this project's keyword list, and flags whether extraction
came back empty (== scanned image needing OCR, see docs/ or
memory/ocr_tooling_setup.md for the tesseract pipeline).

Run:
    .venv312/bin/python3 scripts/264_survey_document_media.py
"""
from __future__ import annotations

import json
import re
import sqlite3
import sys
from pathlib import Path

ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ROOT / "src"))

STATE_DB = ROOT / "data" / "state.sqlite"
RAW_DIR = ROOT / "data" / "raw"
OUT_PATH = ROOT / "data" / "parsed" / "document_media_survey.jsonl"

KEYWORDS = [
    "Мариуполь", "Мариуполя", "Мариуполю",
    "бесхозя", "изъят", "снос", "аварийны", "маневренн",
    "земельного участка", "инвестиционного проекта", "без проведения торгов",
    "муниципальной собственности", "государственной собственности",
    "ипотек", "многоквартирн", "жилищн", "выселен", "компенсаци",
    "ЕГРН", "кадастров", "инвентаризац", "незавершенного строительства",
    "квартир", "дом снес", "переселен", "расселен",
    "распоряжение", "указ главы", "постановление", "приказ",
    "перечень объектов", "перечень жилых", "коммерческих объектов",
    "промышленных площадок", "признаки бесхозности", "признаками бесхозяйного",
]
KEYWORD_RE = re.compile("|".join(re.escape(k) for k in KEYWORDS))


def _extract_pdf(path: Path) -> str:
    import pdfplumber
    text_parts = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages[:30]:  # cap — huge multi-hundred-page PDFs get a representative sample
            t = page.extract_text()
            if t:
                text_parts.append(t)
    return "\n".join(text_parts)


def _extract_docx(path: Path) -> str:
    import docx
    d = docx.Document(str(path))
    parts = [p.text for p in d.paragraphs if p.text]
    for table in d.tables:
        for row in table.rows:
            parts.append(" | ".join(c.text for c in row.cells))
    return "\n".join(parts)


def _extract_xlsx(path: Path) -> str:
    import openpyxl
    wb = openpyxl.load_workbook(path, data_only=True, read_only=True)
    parts = []
    for ws in wb.worksheets:
        for i, row in enumerate(ws.iter_rows(values_only=True)):
            if i > 2000:
                break
            parts.append(" | ".join(str(c) for c in row if c is not None))
    return "\n".join(parts)


def extract_text(path: Path, mime: str) -> tuple[str, str | None]:
    """Returns (text, error). error is set for legacy/unsupported formats."""
    try:
        if mime == "application/pdf":
            return _extract_pdf(path), None
        if mime == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            return _extract_docx(path), None
        if mime == "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet":
            return _extract_xlsx(path), None
        if mime in ("application/msword", "application/vnd.ms-excel"):
            return "", "legacy_binary_office_format_not_extracted"
        return "", f"unhandled_mime:{mime}"
    except Exception as e:  # noqa: BLE001
        return "", f"extract_error:{e}"


def main() -> None:
    con = sqlite3.connect(STATE_DB)
    rows = con.execute(
        "SELECT sha256, url, title, description FROM source_document "
        "WHERE source_type='telegram_document_media' ORDER BY url"
    ).fetchall()
    print(f"{len(rows)} document attachments to survey", file=sys.stderr)

    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    n_hits = 0
    n_empty_needs_ocr = 0
    n_legacy = 0
    n_errors = 0
    with open(OUT_PATH, "w", encoding="utf-8") as out:
        for i, (sha, url, title, desc) in enumerate(rows, 1):
            if i % 50 == 0:
                print(f"  {i}/{len(rows)}...", file=sys.stderr)
            m = re.match(r"https://t\.me/([^/]+)/(\d+)", url)
            channel = m.group(1) if m else "?"
            mime_m = re.search(r"\((.*?), (.*?)\)", desc or "")
            mime = mime_m.group(2) if mime_m else ""
            ext = {
                "application/pdf": ".pdf",
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document": ".docx",
                "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet": ".xlsx",
                "application/msword": ".doc",
                "application/vnd.ms-excel": ".xls",
            }.get(mime, "")
            path = RAW_DIR / f"{sha}{ext}"
            if not path.exists():
                out.write(json.dumps({"sha256": sha, "url": url, "error": "raw_file_missing"}, ensure_ascii=False) + "\n")
                n_errors += 1
                continue

            text, error = extract_text(path, mime)
            record = {
                "sha256": sha, "url": url, "channel": channel, "mime": mime,
                "title": title, "text_len": len(text),
            }
            if error == "legacy_binary_office_format_not_extracted":
                n_legacy += 1
                record["error"] = error
            elif error:
                n_errors += 1
                record["error"] = error
            elif len(text.strip()) < 30:
                n_empty_needs_ocr += 1
                record["needs_ocr"] = True
            else:
                matched_kw = sorted(set(KEYWORD_RE.findall(text)))
                record["keywords"] = matched_kw
                record["first_500"] = text[:500]
                if matched_kw:
                    n_hits += 1
            out.write(json.dumps(record, ensure_ascii=False) + "\n")

    print(f"\nDone. {n_hits} keyword-hit docs, {n_empty_needs_ocr} empty/scanned "
          f"(need OCR), {n_legacy} legacy .doc/.xls (need manual/OCR conversion), "
          f"{n_errors} errors. Index: {OUT_PATH}", file=sys.stderr)


if __name__ == "__main__":
    main()
