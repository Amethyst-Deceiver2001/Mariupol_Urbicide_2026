#!/usr/bin/env python3
"""Convert the 13 legacy .doc/.xls attachments (flagged
error=legacy_binary_office_format_not_extracted by scripts/264) to modern
.docx/.xlsx via LibreOffice headless, then extract their text the same way
scripts/264 does for native .docx/.xlsx.

Requires LibreOffice (`brew install --cask libreoffice` on macOS — it ships
the `soffice` binary used here).

Local, offline, no network.

Output: data/parsed/legacy_converted/<sha256>.{docx,xlsx} (the converted
file) + data/parsed/legacy_office_survey.jsonl (same shape as
document_media_survey.jsonl rows, so scripts/264's downstream keyword-triage
consumers can just concatenate the two files).

Run:
    .venv312/bin/python3 scripts/267_convert_legacy_office.py
"""
from __future__ import annotations

import json
import re
import subprocess
import sys
from pathlib import Path

ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ROOT / "src"))

SURVEY_PATH = ROOT / "data" / "parsed" / "document_media_survey.jsonl"
OUT_DIR = ROOT / "data" / "parsed" / "legacy_converted"
OUT_SURVEY_PATH = ROOT / "data" / "parsed" / "legacy_office_survey.jsonl"
RAW_DIR = ROOT / "data" / "raw"

KEYWORDS = [
    "Мариуполь", "Мариуполя", "Мариуполю",
    "бесхозя", "изъят", "снос", "аварийны", "маневренн",
    "земельного участка", "инвестиционного проекта", "без проведения торгов",
    "муниципальной собственности", "государственной собственности",
    "ипотек", "многоквартирн", "жилищн", "выселен", "компенсаци",
    "ЕГРН", "кадастров", "инвентаризац", "незавершенного строительства",
    "квартир", "дом снес", "переселен", "расселен",
    "распоряжение", "указ главы", "постановление", "приказ",
    "перечень объектов", "перечень жилых", "коммерческих объектов",
    "промышленных площадок", "признаки бесхозности", "признаками бесхозяйного",
    "выморочн", "наследств", "нотариус",
    "росреестр", "недружественн", "иностранн", "украинск",
    "блокированной застройки", "списан", "балансов",
    "спецразрешен", "коллегиальн",
]
KEYWORD_RE = re.compile("|".join(re.escape(k) for k in KEYWORDS))

MIME_EXT = {
    "application/msword": (".doc", "docx", "MS Word 2007 XML"),
    "application/vnd.ms-excel": (".xls", "xlsx", "Calc MS Excel 2007 XML"),
}


def _extract_docx(path: Path) -> str:
    import docx
    d = docx.Document(str(path))
    parts = [p.text for p in d.paragraphs if p.text]
    for table in d.tables:
        for row in table.rows:
            parts.append(" | ".join(c.text for c in row.cells))
    return "\n".join(parts)


def _extract_xlsx(path: Path) -> str:
    import openpyxl
    wb = openpyxl.load_workbook(path, data_only=True, read_only=True)
    parts = []
    for ws in wb.worksheets:
        for i, row in enumerate(ws.iter_rows(values_only=True)):
            if i > 2000:
                break
            parts.append(" | ".join(str(c) for c in row if c is not None))
    return "\n".join(parts)


def main() -> None:
    rows = []
    decoder = json.JSONDecoder()
    text = SURVEY_PATH.read_text(encoding="utf-8")
    pos = 0
    length = len(text)
    while pos < length:
        while pos < length and text[pos] in "\n\r\t ":
            pos += 1
        if pos >= length:
            break
        obj, end = decoder.raw_decode(text, pos)
        rows.append(obj)
        pos = end
    legacy = [r for r in rows if r.get("error") == "legacy_binary_office_format_not_extracted"]
    print(f"{len(legacy)} legacy .doc/.xls documents to convert", file=sys.stderr)

    OUT_DIR.mkdir(parents=True, exist_ok=True)
    n_done = n_errors = 0
    with open(OUT_SURVEY_PATH, "w", encoding="utf-8") as out:
        for i, r in enumerate(legacy, 1):
            sha = r["sha256"]
            mime = r.get("mime", "")
            if mime not in MIME_EXT:
                print(f"  [{i}/{len(legacy)}] unrecognized mime {mime}, skipping", file=sys.stderr)
                n_errors += 1
                continue
            src_ext, target_ext, target_filter = MIME_EXT[mime]
            src_path = RAW_DIR / f"{sha}{src_ext}"
            if not src_path.exists():
                print(f"  [{i}/{len(legacy)}] MISSING raw file: {src_path}", file=sys.stderr)
                n_errors += 1
                continue

            converted_path = OUT_DIR / f"{sha}.{target_ext}"
            print(f"  [{i}/{len(legacy)}] converting {r['url']} ({r.get('title', sha)})", file=sys.stderr)
            try:
                subprocess.run(
                    ["soffice", "--headless", "--convert-to", f"{target_ext}:{target_filter}",
                     "--outdir", str(OUT_DIR), str(src_path)],
                    check=True, capture_output=True, text=True, timeout=60,
                )
            except Exception as e:  # noqa: BLE001
                print(f"    conversion error: {e}", file=sys.stderr)
                n_errors += 1
                continue

            # soffice names the output after the source stem, not the sha — rename to match convention
            produced = OUT_DIR / f"{src_path.stem}.{target_ext}"
            if produced != converted_path and produced.exists():
                produced.rename(converted_path)
            if not converted_path.exists():
                print(f"    conversion produced no output file", file=sys.stderr)
                n_errors += 1
                continue

            try:
                text = _extract_docx(converted_path) if target_ext == "docx" else _extract_xlsx(converted_path)
            except Exception as e:  # noqa: BLE001
                print(f"    extract error: {e}", file=sys.stderr)
                n_errors += 1
                continue

            hits = sorted(set(KEYWORD_RE.findall(text)))
            out.write(json.dumps({
                "sha256": sha, "url": r["url"], "channel": r.get("channel"),
                "title": r.get("title"), "mime": mime, "text_len": len(text),
                "keyword_hits": hits, "converted_path": str(converted_path),
            }, ensure_ascii=False) + "\n")
            n_done += 1

    print(f"\nDone. {n_done} converted+extracted, {n_errors} errors. "
          f"Output: {OUT_SURVEY_PATH}", file=sys.stderr)


if __name__ == "__main__":
    main()
